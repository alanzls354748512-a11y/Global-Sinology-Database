import os, json, requests
from bs4 import BeautifulSoup
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document

# --- 1. 核心配置：路徑與 ID 精確映射 ---
FOLDER_MAP = {
    'Document_Geography': '12Y0tfBUQ-B6VZPEVTLIFKlAleY9GIDSa',
    'East_Asian_History': '1409gDpMZT0Ew3-J2t6Sbr-6BffZH4gZ4',
    'Thought_Governance': '14H9f4hduc3QmmE3TAjnCtvNn36xdVHJU',
    'Cross_Analysis': '1BxkNCkitbw-YMO0BDcQzdOG6KmXEXR0W'
}

# --- 2. 模擬學術爬蟲邏輯 ---
def fetch_academic_content(keyword):
    """
    未來可在此接入 Google Scholar 或 CrossRef API
    目前先以關鍵字自動生成具備 NSS 邏輯的分析摘要
    """
    report_content = f"""
    【全球漢學監測報表 - 關鍵字：{keyword}】
    
    1. 最新研究趨勢：
    當前學界針對 {keyword} 的討論已從單一史料轉向跨國地緣政治對標。
    
    2. 與 NSS (美國國家安全戰略) 的聯繫：
    透過歷史週期看現代供應鏈韌性，{keyword} 的治理經驗提供了重要的政權穩定性參考。
    
    3. 數據庫建議：
    建議將此資產納入「全球金融數據庫」的地緣政治風險分析模型中。
    """
    return report_content

def upload_as_docx(text_content, title, category):
    creds_json = os.environ.get('GDRIVE_CREDENTIALS')
    if not creds_json:
        print("❌ 錯誤：未找到 GDRIVE_CREDENTIALS 密鑰")
        return

    creds_dict = json.loads(creds_json)
    creds = service_account.Credentials.from_service_account_info(creds_dict)
    service = build('drive', 'v3', credentials=creds)

    # 建立 Word 文檔
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text_content)
    
    filename = f"{title}.docx"
    doc.save(filename)

    # 獲取資料夾 ID
    folder_id = FOLDER_MAP.get(category)
    if not folder_id:
        print(f"⚠️ 找不到類別 {category} 對應的 ID")
        return

    # 上傳至 Google Drive
    file_metadata = {'name': filename, 'parents': [folder_id]}
    media = MediaFileUpload(filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    try:
        service.files().create(body=file_metadata, media_body=media).execute()
        print(f"✅ 已成功歸檔至 {category}: {filename}")
    except Exception as e:
        print(f"❌ 歸檔失敗: {str(e)}")

# --- 3. 自動化運行清單 ---
if __name__ == "__main__":
    # 這是您每天自動監控的學術任務清單
    TASKS = [
        {"kw": "清代行政效率與糧食安全", "cat": "Thought_Governance"},
        {"kw": "東亞海域地緣政治與白銀本位", "cat": "East_Asian_History"},
        {"kw": "近代中日金融條約與主權變遷", "cat": "Cross_Analysis"}
    ]
    
    for task in TASKS:
        content = fetch_academic_content(task['kw'])
        upload_as_docx(content, f"自動監測_{task['kw']}", task['cat'])
